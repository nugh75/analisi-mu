"""
Blueprint per la gestione del Diario di Bordo con funzionalità di esportazione
"""

from flask import Blueprint, render_template, request, redirect, url_for, flash, jsonify, current_app, send_file
from flask_login import login_required, current_user
from sqlalchemy import desc, and_, or_
from models import db, DiaryEntry, DiaryAttachment, User, ExcelFile
from datetime import datetime, timedelta
import os
import json
import re
from werkzeug.utils import secure_filename
import tempfile
import zipfile
from io import BytesIO

# Importazioni per export
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.pdfgen import canvas
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

diary_bp = Blueprint('diary', __name__)

ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif', 'pdf', 'doc', 'docx', 'txt', 'xlsx', 'csv'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@diary_bp.route('/')
@login_required
def index():
    """Pagina principale del diario di bordo"""
    page = request.args.get('page', 1, type=int)
    activity_type = request.args.get('type', 'all')
    priority = request.args.get('priority', 'all')
    status = request.args.get('status', 'all')
    search = request.args.get('search', '')
    
    # Costruisci la query base
    query = DiaryEntry.query
    
    # Applica filtri
    if activity_type != 'all':
        query = query.filter(DiaryEntry.activity_type == activity_type)
    
    if priority != 'all':
        query = query.filter(DiaryEntry.priority == priority)
    
    if status != 'all':
        query = query.filter(DiaryEntry.status == status)
    
    if search:
        search_term = f"%{search}%"
        query = query.filter(
            or_(
                DiaryEntry.title.ilike(search_term),
                DiaryEntry.content.ilike(search_term),
                DiaryEntry.tags.ilike(search_term)
            )
        )
    
    # Ordina per data (più recenti primi)
    query = query.order_by(desc(DiaryEntry.created_at))
    
    # Paginazione
    entries = query.paginate(
        page=page, 
        per_page=10, 
        error_out=False
    )
    
    # Statistiche per la dashboard
    stats = {
        'total': DiaryEntry.query.count(),
        'active': DiaryEntry.query.filter_by(status='active').count(),
        'completed': DiaryEntry.query.filter_by(status='completed').count(),
        'high_priority': DiaryEntry.query.filter_by(priority='high').count(),
        'urgent': DiaryEntry.query.filter_by(priority='urgent').count()
    }
    
    # Progetti disponibili
    projects = ExcelFile.query.all()
    
    return render_template('diary/index.html', 
                         entries=entries, 
                         stats=stats,
                         projects=projects,
                         current_filters={
                             'type': activity_type,
                             'priority': priority,
                             'status': status,
                             'search': search
                         })

@diary_bp.route('/create', methods=['GET', 'POST'])
@login_required
def create():
    """Crea una nuova voce del diario"""
    if request.method == 'POST':
        title = request.form.get('title')
        content = request.form.get('content')
        activity_type = request.form.get('activity_type', 'general')
        priority = request.form.get('priority', 'medium')
        project_id = request.form.get('project_id')
        tags = request.form.get('tags', '')
        entry_date_str = request.form.get('entry_date')
        
        if not title or not content:
            flash('Titolo e contenuto sono obbligatori', 'error')
            return redirect(url_for('diary.create'))
        
        # Gestisci la data della voce
        entry_date = datetime.utcnow()
        if entry_date_str:
            try:
                entry_date = datetime.strptime(entry_date_str, '%Y-%m-%dT%H:%M')
            except ValueError:
                try:
                    # Prova formato solo data
                    entry_date = datetime.strptime(entry_date_str, '%Y-%m-%d')
                except ValueError:
                    entry_date = datetime.utcnow()
        
        # Crea nuova voce
        entry = DiaryEntry(
            title=title,
            content=content,
            activity_type=activity_type,
            priority=priority,
            author_id=current_user.id,
            project_id=project_id if project_id else None,
            entry_date=entry_date
        )
        
        # Elabora tag
        if tags:
            tags_list = [tag.strip() for tag in tags.split(',') if tag.strip()]
            entry.tags_list = tags_list
        
        # Analizza menzioni e riferimenti
        entry.parse_mentions_and_files()
        entry.update_word_count()
        
        db.session.add(entry)
        db.session.commit()
        
        # Gestisci eventuali allegati
        if 'attachments' in request.files:
            files = request.files.getlist('attachments')
            for file in files:
                if file and file.filename and allowed_file(file.filename):
                    filename = secure_filename(file.filename)
                    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                    unique_filename = f"{timestamp}_{filename}"
                    
                    # Crea directory se non exists
                    upload_dir = os.path.join(current_app.config.get('UPLOAD_FOLDER', 'uploads'), 'diary')
                    os.makedirs(upload_dir, exist_ok=True)
                    
                    file_path = os.path.join(upload_dir, unique_filename)
                    file.save(file_path)
                    
                    # Crea record allegato
                    attachment = DiaryAttachment(
                        entry_id=entry.id,
                        filename=unique_filename,
                        original_name=filename,
                        file_path=file_path,
                        file_size=os.path.getsize(file_path),
                        mime_type=file.content_type
                    )
                    db.session.add(attachment)
        
        db.session.commit()
        flash('Voce del diario creata con successo', 'success')
        return redirect(url_for('diary.view', id=entry.id))
    
    # GET - mostra form
    projects = ExcelFile.query.all()
    return render_template('diary/create.html', projects=projects)

@diary_bp.route('/view/<int:id>')
@login_required
def view(id):
    """Visualizza una voce del diario"""
    entry = DiaryEntry.query.get_or_404(id)
    
    # Elabora menzioni per creare link
    content_with_links = process_content_links(entry.content)
    
    return render_template('diary/view.html', 
                         entry=entry, 
                         content_with_links=content_with_links)

@diary_bp.route('/edit/<int:id>', methods=['GET', 'POST'])
@login_required
def edit(id):
    """Modifica una voce del diario"""
    entry = DiaryEntry.query.get_or_404(id)
    
    # Verifica permessi
    if entry.author_id != current_user.id and not current_user.is_admin:
        flash('Non hai i permessi per modificare questa voce', 'error')
        return redirect(url_for('diary.view', id=id))
    
    if request.method == 'POST':
        entry.title = request.form.get('title')
        entry.content = request.form.get('content')
        entry.activity_type = request.form.get('activity_type', 'general')
        entry.priority = request.form.get('priority', 'medium')
        entry.status = request.form.get('status', 'active')
        entry.project_id = request.form.get('project_id') if request.form.get('project_id') else None
        
        # Gestisci aggiornamento data
        entry_date_str = request.form.get('entry_date')
        if entry_date_str:
            try:
                entry.entry_date = datetime.strptime(entry_date_str, '%Y-%m-%dT%H:%M')
            except ValueError:
                try:
                    entry.entry_date = datetime.strptime(entry_date_str, '%Y-%m-%d')
                except ValueError:
                    pass  # Mantieni la data esistente se il formato non è valido
        
        # Aggiorna tag
        tags = request.form.get('tags', '')
        if tags:
            tags_list = [tag.strip() for tag in tags.split(',') if tag.strip()]
            entry.tags_list = tags_list
        else:
            entry.tags_list = []
        
        # Rianalizza menzioni e riferimenti
        entry.parse_mentions_and_files()
        entry.update_word_count()
        entry.updated_at = datetime.utcnow()
        
        db.session.commit()
        flash('Voce del diario aggiornata con successo', 'success')
        return redirect(url_for('diary.view', id=id))
    
    # GET - mostra form di modifica
    projects = ExcelFile.query.all()
    return render_template('diary/edit.html', entry=entry, projects=projects)

@diary_bp.route('/delete/<int:id>', methods=['POST'])
@login_required
def delete(id):
    """Elimina una voce del diario"""
    entry = DiaryEntry.query.get_or_404(id)
    
    # Verifica permessi
    if entry.author_id != current_user.id and not current_user.is_admin:
        flash('Non hai i permessi per eliminare questa voce', 'error')
        return redirect(url_for('diary.view', id=id))
    
    # Elimina allegati fisici
    for attachment in entry.attachments:
        try:
            if os.path.exists(attachment.file_path):
                os.remove(attachment.file_path)
        except:
            pass
    
    db.session.delete(entry)
    db.session.commit()
    
    flash('Voce del diario eliminata con successo', 'success')
    return redirect(url_for('diary.index'))

@diary_bp.route('/export')
@login_required
def export_options():
    """Pagina per le opzioni di esportazione"""
    return render_template('diary/export.html')

@diary_bp.route('/export/txt')
@login_required
def export_txt():
    """Esporta il diario in formato TXT"""
    # Parametri di filtro
    start_date = request.args.get('start_date')
    end_date = request.args.get('end_date')
    activity_type = request.args.get('activity_type', 'all')
    priority = request.args.get('priority', 'all')
    status = request.args.get('status', 'all')
    
    # Costruisci query con filtri
    query = build_export_query(start_date, end_date, activity_type, priority, status)
    entries = query.all()
    
    # Genera contenuto TXT
    content = generate_txt_content(entries)
    
    # Crea file temporaneo
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    filename = f"diario_bordo_{timestamp}.txt"
    
    temp_file = tempfile.NamedTemporaryFile(mode='w', encoding='utf-8', delete=False, suffix='.txt')
    temp_file.write(content)
    temp_file.close()
    
    return send_file(temp_file.name, as_attachment=True, download_name=filename, mimetype='text/plain')

@diary_bp.route('/export/word')
@login_required
def export_word():
    """Esporta il diario in formato Word"""
    if not DOCX_AVAILABLE:
        flash('Libreria python-docx non disponibile per export Word', 'error')
        return redirect(url_for('diary.export_options'))
    
    # Parametri di filtro
    start_date = request.args.get('start_date')
    end_date = request.args.get('end_date')
    activity_type = request.args.get('activity_type', 'all')
    priority = request.args.get('priority', 'all')
    status = request.args.get('status', 'all')
    
    # Costruisci query con filtri
    query = build_export_query(start_date, end_date, activity_type, priority, status)
    entries = query.all()
    
    # Crea documento Word
    doc = generate_word_document(entries)
    
    # Salva in file temporaneo
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    filename = f"diario_bordo_{timestamp}.docx"
    
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)
    temp_file.close()
    
    return send_file(temp_file.name, as_attachment=True, download_name=filename, 
                    mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

@diary_bp.route('/export/pdf')
@login_required
def export_pdf():
    """Esporta il diario in formato PDF"""
    if not PDF_AVAILABLE:
        flash('Libreria reportlab non disponibile per export PDF', 'error')
        return redirect(url_for('diary.export_options'))
    
    # Parametri di filtro
    start_date = request.args.get('start_date')
    end_date = request.args.get('end_date')
    activity_type = request.args.get('activity_type', 'all')
    priority = request.args.get('priority', 'all')
    status = request.args.get('status', 'all')
    
    # Costruisci query con filtri
    query = build_export_query(start_date, end_date, activity_type, priority, status)
    entries = query.all()
    
    # Crea documento PDF
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    filename = f"diario_bordo_{timestamp}.pdf"
    
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
    generate_pdf_document(entries, temp_file.name)
    temp_file.close()
    
    return send_file(temp_file.name, as_attachment=True, download_name=filename, mimetype='application/pdf')

@diary_bp.route('/api/users')
@login_required
def api_users():
    """API per autocompletamento utenti"""
    search = request.args.get('q', '')
    if len(search) < 2:
        return jsonify([])
    
    users = User.query.filter(
        User.username.ilike(f"%{search}%")
    ).limit(10).all()
    
    return jsonify([{
        'id': user.id,
        'username': user.username,
        'display_name': f"{user.username} ({user.email})"
    } for user in users])

@diary_bp.route('/api/files')
@login_required
def api_files():
    """API per autocompletamento file"""
    search = request.args.get('q', '')
    if len(search) < 2:
        return jsonify([])
    
    # Cerca nei file Excel caricati
    files = ExcelFile.query.filter(
        ExcelFile.original_filename.ilike(f"%{search}%")
    ).limit(10).all()
    
    results = []
    for file in files:
        results.append({
            'name': file.original_filename,
            'type': 'excel',
            'path': file.file_path
        })
    
    return jsonify(results)

# Funzioni di supporto

def get_activity_type_display(activity_type):
    """Traduce il tipo di attività in italiano"""
    translations = {
        'general': 'Generale',
        'meeting': 'Riunione',
        'milestone': 'Milestone',
        'issue': 'Problema',
        'decision': 'Decisione',
        'reflection': 'Riflessione',
        'consideration': 'Considerazione'
    }
    return translations.get(activity_type, activity_type.title())

def build_export_query(start_date, end_date, activity_type, priority, status):
    """Costruisce la query per l'export con filtri"""
    query = DiaryEntry.query
    
    # Filtro per date
    if start_date:
        try:
            start = datetime.strptime(start_date, '%Y-%m-%d')
            query = query.filter(DiaryEntry.created_at >= start)
        except ValueError:
            pass
    
    if end_date:
        try:
            end = datetime.strptime(end_date, '%Y-%m-%d')
            # Aggiungi 23:59:59 per includere tutto il giorno
            end = end.replace(hour=23, minute=59, second=59)
            query = query.filter(DiaryEntry.created_at <= end)
        except ValueError:
            pass
    
    # Altri filtri
    if activity_type != 'all':
        query = query.filter(DiaryEntry.activity_type == activity_type)
    
    if priority != 'all':
        query = query.filter(DiaryEntry.priority == priority)
    
    if status != 'all':
        query = query.filter(DiaryEntry.status == status)
    
    return query.order_by(desc(DiaryEntry.created_at))

def generate_txt_content(entries):
    """Genera il contenuto del file TXT"""
    content = []
    content.append("DIARIO DI BORDO")
    content.append("=" * 50)
    content.append(f"Esportato il: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    content.append(f"Numero di voci: {len(entries)}")
    content.append("")
    
    for entry in entries:
        content.append("-" * 50)
        content.append(f"TITOLO: {entry.title}")
        content.append(f"AUTORE: {entry.author.username}")
        content.append(f"DATA: {entry.created_at.strftime('%d/%m/%Y %H:%M')}")
        content.append(f"TIPO: {get_activity_type_display(entry.activity_type).upper()}")
        content.append(f"PRIORITÀ: {entry.priority.upper()}")
        content.append(f"STATUS: {entry.status.upper()}")
        
        if entry.project:
            content.append(f"PROGETTO: {entry.project.original_filename}")
        
        if entry.tags_list:
            content.append(f"TAG: {', '.join(entry.tags_list)}")
        
        if entry.mentioned_users_list:
            content.append(f"UTENTI MENZIONATI: {', '.join(['@' + u for u in entry.mentioned_users_list])}")
        
        if entry.referenced_files_list:
            content.append(f"FILE REFERENZIATI: {', '.join(['#' + f for f in entry.referenced_files_list])}")
        
        content.append("")
        content.append("CONTENUTO:")
        content.append(entry.content)
        content.append("")
        
        if entry.attachments:
            content.append("ALLEGATI:")
            for att in entry.attachments:
                content.append(f"  - {att.original_name} ({att.file_size_human})")
            content.append("")
    
    return "\n".join(content)

def generate_word_document(entries):
    """Genera il documento Word"""
    doc = Document()
    
    # Titolo principale
    title = doc.add_heading('Diario di Bordo', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Info export
    info_para = doc.add_paragraph()
    info_para.add_run(f"Esportato il: ").bold = True
    info_para.add_run(datetime.now().strftime('%d/%m/%Y %H:%M'))
    info_para.add_run(f"\nNumero di voci: ").bold = True
    info_para.add_run(str(len(entries)))
    
    doc.add_page_break()
    
    for i, entry in enumerate(entries):
        # Titolo voce
        heading = doc.add_heading(entry.title, level=1)
        
        # Tabella con metadati
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        metadata = [
            ('Autore', entry.author.username),
            ('Data', entry.created_at.strftime('%d/%m/%Y %H:%M')),
            ('Tipo', get_activity_type_display(entry.activity_type)),
            ('Priorità', entry.priority.title()),
            ('Status', entry.status.title())
        ]
        
        if entry.project:
            metadata.append(('Progetto', entry.project.original_filename))
        
        if entry.tags_list:
            metadata.append(('Tag', ', '.join(entry.tags_list)))
        
        if entry.mentioned_users_list:
            metadata.append(('Utenti menzionati', ', '.join(['@' + u for u in entry.mentioned_users_list])))
        
        if entry.referenced_files_list:
            metadata.append(('File referenziati', ', '.join(['#' + f for f in entry.referenced_files_list])))
        
        # Aggiungi righe alla tabella
        for key, value in metadata:
            row_cells = table.add_row().cells
            row_cells[0].text = key
            row_cells[1].text = value
            row_cells[0].paragraphs[0].runs[0].bold = True
        
        # Contenuto
        doc.add_heading('Contenuto', level=2)
        content_para = doc.add_paragraph(entry.content)
        
        # Allegati
        if entry.attachments:
            doc.add_heading('Allegati', level=2)
            for att in entry.attachments:
                att_para = doc.add_paragraph()
                att_para.add_run(f"📎 {att.original_name}").bold = True
                att_para.add_run(f" ({att.file_size_human})")
        
        # Separatore se non è l'ultima voce
        if i < len(entries) - 1:
            doc.add_page_break()
    
    return doc

def generate_pdf_document(entries, output_path):
    """Genera il documento PDF"""
    doc = SimpleDocTemplate(output_path, pagesize=A4)
    styles = getSampleStyleSheet()
    story = []
    
    # Stile personalizzato per il titolo
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        spaceAfter=30,
        alignment=TA_CENTER
    )
    
    # Titolo principale
    story.append(Paragraph("Diario di Bordo", title_style))
    story.append(Spacer(1, 20))
    
    # Info export
    info_text = f"""
    <b>Esportato il:</b> {datetime.now().strftime('%d/%m/%Y %H:%M')}<br/>
    <b>Numero di voci:</b> {len(entries)}
    """
    story.append(Paragraph(info_text, styles['Normal']))
    story.append(Spacer(1, 30))
    
    for i, entry in enumerate(entries):
        # Titolo voce
        story.append(Paragraph(entry.title, styles['Heading1']))
        story.append(Spacer(1, 12))
        
        # Metadati in tabella
        metadata_data = [
            ['<b>Campo</b>', '<b>Valore</b>'],
            ['Autore', entry.author.username],
            ['Data', entry.created_at.strftime('%d/%m/%Y %H:%M')],
            ['Tipo', get_activity_type_display(entry.activity_type)],
            ['Priorità', entry.priority.title()],
            ['Status', entry.status.title()]
        ]
        
        if entry.project:
            metadata_data.append(['Progetto', entry.project.original_filename])
        
        if entry.tags_list:
            metadata_data.append(['Tag', ', '.join(entry.tags_list)])
        
        if entry.mentioned_users_list:
            metadata_data.append(['Utenti menzionati', ', '.join(['@' + u for u in entry.mentioned_users_list])])
        
        if entry.referenced_files_list:
            metadata_data.append(['File referenziati', ', '.join(['#' + f for f in entry.referenced_files_list])])
        
        metadata_table = Table(metadata_data, colWidths=[2*inch, 4*inch])
        metadata_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        story.append(metadata_table)
        story.append(Spacer(1, 20))
        
        # Contenuto
        story.append(Paragraph("<b>Contenuto:</b>", styles['Heading2']))
        
        # Divide il contenuto in paragrafi
        content_paragraphs = entry.content.split('\n')
        for para in content_paragraphs:
            if para.strip():
                story.append(Paragraph(para, styles['Normal']))
                story.append(Spacer(1, 6))
        
        # Allegati
        if entry.attachments:
            story.append(Spacer(1, 12))
            story.append(Paragraph("<b>Allegati:</b>", styles['Heading2']))
            for att in entry.attachments:
                att_text = f"📎 <b>{att.original_name}</b> ({att.file_size_human})"
                story.append(Paragraph(att_text, styles['Normal']))
        
        # Separatore se non è l'ultima voce
        if i < len(entries) - 1:
            story.append(Spacer(1, 30))
            story.append(Paragraph("<hr/>", styles['Normal']))
            story.append(Spacer(1, 30))
    
    doc.build(story)

def process_content_links(content):
    """Processa il contenuto per creare link per menzioni e riferimenti file"""
    # Prima controlla se il contenuto è già HTML (da Quill)
    if '<p>' in content or '<h1>' in content or '<h2>' in content or '<h3>' in content:
        # È già HTML, processa solo le menzioni
        processed_content = content
    else:
        # È testo semplice o Markdown, processa Markdown e aggiungi a capo
        processed_content = process_markdown_and_newlines(content)
    
    # Sostituisci menzioni utenti
    processed_content = re.sub(
        r'@(\w+)',
        r'<span class="mention-user" title="Utente menzionato">@\1</span>',
        processed_content
    )
    
    # Sostituisci riferimenti file
    processed_content = re.sub(
        r'#(\S+)',
        r'<span class="mention-file" title="File referenziato">#\1</span>',
        processed_content
    )
    
    return processed_content

def process_markdown_and_newlines(content):
    """Processa Markdown base e converte a capo in HTML"""
    # Processa Markdown base
    # Titoli
    content = re.sub(r'^### (.*?)$', r'<h3>\1</h3>', content, flags=re.MULTILINE)
    content = re.sub(r'^## (.*?)$', r'<h2>\1</h2>', content, flags=re.MULTILINE)
    content = re.sub(r'^# (.*?)$', r'<h1>\1</h1>', content, flags=re.MULTILINE)
    
    # Grassetto e corsivo
    content = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', content)
    content = re.sub(r'\*(.*?)\*', r'<em>\1</em>', content)
    
    # Liste puntate
    content = re.sub(r'^[\*\-\+] (.*?)$', r'<li>\1</li>', content, flags=re.MULTILINE)
    content = re.sub(r'(<li>.*?</li>\s*)+', lambda m: f'<ul>{m.group(0)}</ul>', content, flags=re.DOTALL)
    
    # Liste numerate
    content = re.sub(r'^\d+\. (.*?)$', r'<li>\1</li>', content, flags=re.MULTILINE)
    # Sostituisci sequenze di <li> numerati con <ol>
    content = re.sub(r'(<li>.*?</li>(?:\s*<li>.*?</li>)*)', lambda m: f'<ol>{m.group(0)}</ol>' if '.' in content else m.group(0), content)
    
    # Paragrafi (doppio a capo)
    paragraphs = content.split('\n\n')
    processed_paragraphs = []
    
    for para in paragraphs:
        para = para.strip()
        if para:
            # Se non è già un elemento HTML, avvolgilo in <p>
            if not (para.startswith('<') and para.endswith('>')):
                # Sostituisci singoli a capo con <br>
                para = para.replace('\n', '<br>')
                para = f'<p>{para}</p>'
            processed_paragraphs.append(para)
    
    return '\n'.join(processed_paragraphs)
