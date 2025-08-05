"""
Route per la gestione dei documenti di testo (focus group, interviste)
"""

import os
import uuid
from datetime import datetime
from flask import Blueprint, render_template, request, jsonify, redirect, url_for, flash, current_app
from flask_login import login_required, current_user
from werkzeug.utils import secure_filename
from docx import Document
import bleach

from models import db, TextDocument, TextAnnotation, Label, Category
from forms import TextDocumentForm

text_documents_bp = Blueprint('text_documents', __name__, url_prefix='/text-documents')

ALLOWED_EXTENSIONS = {'txt', 'md', 'docx'}
MAX_FILE_SIZE = 5 * 1024 * 1024  # Ridotto a 5MB per performance migliori
MAX_CONTENT_LENGTH = 500000  # 500KB di testo puro (circa 100-150 pagine)

def allowed_file(filename):
    """Verifica se il file è di un tipo consentito"""
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def get_file_format(filename):
    """Estrae il formato del file"""
    return filename.rsplit('.', 1)[1].lower() if '.' in filename else 'unknown'

def read_text_file(file_path, file_format):
    """Legge il contenuto di un file di testo"""
    try:
        if file_format == 'docx':
            doc = Document(file_path)
            # Ottimizzazione: usa list comprehension più efficiente
            # e processa solo paragrafi non vuoti
            paragraphs = []
            for p in doc.paragraphs:
                text = p.text.strip()
                if text:  # Solo paragrafi con contenuto
                    paragraphs.append(text)
            content = '\n'.join(paragraphs)
        else:  # txt, md
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
        
        # Sanitizzazione condizionale: solo se contiene HTML
        if '<' in content and '>' in content:
            content = bleach.clean(content, tags=[], strip=True)
        
        return content
    except Exception as e:
        current_app.logger.error(f"Errore nella lettura del file {file_path}: {str(e)}")
        return None

@text_documents_bp.route('/upload-ajax', methods=['POST'])
@login_required
def upload_ajax():
    """Upload AJAX con feedback in tempo reale"""
    try:
        # Verifica presenza file
        if 'file' not in request.files:
            return jsonify({'error': 'Nessun file selezionato'}), 400
        
        file = request.files['file']
        document_type = request.form.get('document_type', 'other')
        
        if file.filename == '':
            return jsonify({'error': 'Nessun file selezionato'}), 400
        
        if not allowed_file(file.filename):
            return jsonify({'error': 'Formato file non supportato'}), 400
        
        # Step 1: Verifica dimensione file
        file.seek(0, os.SEEK_END)
        file_size = file.tell()
        file.seek(0)
        
        if file_size > MAX_FILE_SIZE:
            return jsonify({'error': f'File troppo grande. Dimensione massima: {MAX_FILE_SIZE//1024//1024}MB'}), 400
        
        # Step 2: Salvataggio file
        original_filename = secure_filename(file.filename)
        file_format = get_file_format(original_filename)
        filename = f"{uuid.uuid4().hex}_{original_filename}"
        
        uploads_dir = current_app.config.get('UPLOAD_FOLDER', 'uploads')
        os.makedirs(uploads_dir, exist_ok=True)
        file_path = os.path.join(uploads_dir, filename)
        
        # Salva file
        file.save(file_path)
        current_app.logger.info(f"File salvato: {filename}")
        
        # Step 3: Lettura contenuto (con logging per debug)
        current_app.logger.info(f"Inizio lettura file: {original_filename} ({file_format})")
        content = read_text_file(file_path, file_format)
        
        if content is None:
            os.remove(file_path)  # Cleanup
            return jsonify({'error': 'Errore nella lettura del file'}), 400
        
        current_app.logger.info(f"Contenuto letto: {len(content)} caratteri")
        
        # Verifica lunghezza contenuto
        if len(content) > MAX_CONTENT_LENGTH:
            os.remove(file_path)  # Cleanup
            return jsonify({'error': f'Contenuto troppo lungo. Massimo: {MAX_CONTENT_LENGTH//1000}KB di testo'}), 400
        
        # Step 4: Salvataggio nel database
        current_app.logger.info(f"Salvataggio nel database: {len(content)} caratteri")
        new_document = TextDocument(
            filename=filename,  # Nome file sul server
            original_name=original_filename,  # Nome originale
            content=content,
            document_type=document_type,
            file_format=file_format,
            user_id=current_user.id  # Corretto nome campo
        )
        
        # Aggiorna statistiche
        new_document.word_count = len(content.split()) if content else 0
        new_document.character_count = len(content) if content else 0
        
        db.session.add(new_document)
        db.session.commit()
        
        current_app.logger.info(f"Upload completato: {original_filename}")
        
        return jsonify({
            'success': True,
            'message': 'File caricato con successo!',
            'document_id': new_document.id,
            'redirect_url': url_for('text_documents.annotate', document_id=new_document.id)
        })
        
    except Exception as e:
        current_app.logger.error(f"Errore upload: {str(e)}")
        # Cleanup file se esiste
        if 'file_path' in locals() and os.path.exists(file_path):
            os.remove(file_path)
        return jsonify({'error': f'Errore interno: {str(e)}'}), 500

@text_documents_bp.route('/upload', methods=['GET', 'POST'])
@login_required
def upload():
    """Upload di un nuovo documento di testo"""
    if request.method == 'POST':
        # Per richieste non-AJAX, usa l'endpoint tradizionale
        return upload_ajax()
    
    # GET: mostra form upload
    return render_template('text_documents/upload.html')

@text_documents_bp.route('/list')
@login_required
def list_documents():
    """Lista dei documenti caricati"""
    page = request.args.get('page', 1, type=int)
    per_page = 20
    
    # Filtri
    document_type = request.args.get('type')
    search = request.args.get('search')
    
    query = TextDocument.query
    
    # Applica filtri
    if document_type and document_type != 'all':
        query = query.filter_by(document_type=document_type)
    
    if search:
        search_term = f"%{search}%"
        query = query.filter(TextDocument.original_name.ilike(search_term))
    
    # Se non admin, mostra solo i propri documenti
    if not current_user.is_admin:
        query = query.filter_by(user_id=current_user.id)
    
    documents = query.order_by(TextDocument.created_at.desc()).paginate(
        page=page, per_page=per_page, error_out=False
    )
    
    # Statistiche
    if current_user.is_admin:
        total_query = TextDocument.query
    else:
        total_query = TextDocument.query.filter_by(user_id=current_user.id)
    
    stats = {
        'total': total_query.count(),
        'focus_groups': total_query.filter_by(document_type='focus_group').count(),
        'interviews': total_query.filter_by(document_type='interview').count(),
        'other': total_query.filter_by(document_type='other').count()
    }
    
    return render_template('text_documents/list.html', 
                         documents=documents, 
                         stats=stats,
                         current_type=document_type,
                         current_search=search)

@text_documents_bp.route('/annotate/<int:document_id>')
@login_required
def annotate(document_id):
    """Visualizza un documento per l'etichettamento"""
    document = TextDocument.query.get_or_404(document_id)
    
    # Verifica permessi
    if not current_user.is_admin and document.user_id != current_user.id:
        flash('Non hai i permessi per visualizzare questo documento', 'error')
        return redirect(url_for('text_documents.list_documents'))
    
    # Carica etichette disponibili con le loro categorie
    labels = db.session.query(Label)\
                      .options(db.joinedload(Label.category_obj))\
                      .filter_by(is_active=True)\
                      .order_by(Label.name).all()
    
    # Carica annotazioni esistenti con dettagli utente e etichetta
    annotations = db.session.query(TextAnnotation)\
                           .options(db.joinedload(TextAnnotation.user))\
                           .options(db.joinedload(TextAnnotation.label).joinedload(Label.category_obj))\
                           .filter(TextAnnotation.document_id == document_id)\
                           .order_by(TextAnnotation.start_position).all()
    
    # Statistiche documento
    stats = {
        'annotations_count': len(annotations),
        'unique_labels': len({ann.label_id for ann in annotations}),
        'annotators': len({ann.user_id for ann in annotations}),
        'coverage_percentage': 0
    }
    
    # Calcola copertura (caratteri etichettati vs totali)
    if document.content and len(document.content) > 0:
        annotated_chars = sum(ann.end_position - ann.start_position for ann in annotations)
        stats['coverage_percentage'] = min(100, (annotated_chars / len(document.content)) * 100)
    
    # Converti labels e annotations in formato JSON-serializzabile
    labels_json = [
        {
            'id': label.id,
            'name': label.name,
            'description': label.description,
            'category': label.category_obj.name if label.category_obj else None,
            'color': label.get_effective_color()
        }
        for label in labels
    ]
    
    annotations_json = [
        {
            'id': ann.id,
            'start_position': ann.start_position,
            'end_position': ann.end_position,
            'text_selection': ann.text_selection,
            'label_id': ann.label_id,
            'label_name': ann.label.name if ann.label else f'Etichetta {ann.label_id}',
            'label_color': ann.label.get_effective_color() if ann.label else '#007bff',
            'label_category': ann.label.category_obj.name if ann.label and ann.label.category_obj else None,
            'user_id': ann.user_id,
            'user_name': ann.user.username if ann.user else 'Utente sconosciuto',
            'created_at': ann.created_at.strftime('%d/%m/%Y %H:%M') if ann.created_at else ''
        }
        for ann in annotations
    ]
    
    return render_template('text_documents/annotate.html',
                         document=document,
                         labels_json=labels_json,
                         annotations_json=annotations_json,
                         stats=stats)

@text_documents_bp.route('/api/annotate', methods=['POST'])
@login_required
def api_annotate():
    """API per creare una nuova annotazione"""
    try:
        data = request.get_json()
        
        document_id = data.get('document_id')
        text_selection = data.get('text_selection')
        start_position = data.get('start_position')
        end_position = data.get('end_position')
        label_id = data.get('label_id')
        context_before = data.get('context_before', '')
        context_after = data.get('context_after', '')
        
        # Validazioni
        if not all([document_id, text_selection, start_position is not None, 
                   end_position is not None, label_id]):
            return jsonify({'success': False, 'message': 'Dati mancanti'}), 400
        
        # Verifica documento
        document = TextDocument.query.get(document_id)
        if not document:
            return jsonify({'success': False, 'message': 'Documento non trovato'}), 404
        
        # Verifica permessi
        if not current_user.is_admin and document.user_id != current_user.id:
            return jsonify({'success': False, 'message': 'Permessi insufficienti'}), 403
        
        # Verifica etichetta
        label = Label.query.get(label_id)
        if not label or not label.is_active:
            return jsonify({'success': False, 'message': 'Etichetta non valida'}), 400
        

        
        # Crea annotazione
        annotation = TextAnnotation(
            document_id=document_id,
            text_selection=text_selection,
            start_position=start_position,
            end_position=end_position,
            label_id=label_id,
            user_id=current_user.id,
            context_before=context_before,
            context_after=context_after
        )
        
        db.session.add(annotation)
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': 'Annotazione creata con successo',
            'annotation': {
                'id': annotation.id,
                'start_position': annotation.start_position,
                'end_position': annotation.end_position,
                'text_selection': annotation.text_selection,
                'label_id': annotation.label_id,
                'label_name': label.name,
                'label_color': label.get_effective_color(),
                'label_category': label.category_obj.name if label.category_obj else None,
                'user_id': annotation.user_id,
                'user_name': current_user.username,
                'created_at': annotation.created_at.strftime('%d/%m/%Y %H:%M')
            }
        })
        
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Errore nella creazione annotazione: {str(e)}")
        return jsonify({'success': False, 'message': 'Errore interno del server'}), 500

@text_documents_bp.route('/annotations/<int:document_id>')
@login_required
def annotations(document_id):
    """Mostra tutte le annotazioni di un documento"""
    document = TextDocument.query.get_or_404(document_id)
    
    # Verifica permessi
    if not current_user.is_admin and document.user_id != current_user.id:
        flash('Non hai i permessi per visualizzare questo documento', 'error')
        return redirect(url_for('text_documents.list_documents'))
    
    # Parametri di filtro
    selected_categories = request.args.getlist('categories')
    
    # Query base per le annotazioni
    annotations_query = db.session.query(TextAnnotation)\
                               .options(db.joinedload(TextAnnotation.label).joinedload(Label.category_obj))\
                               .options(db.joinedload(TextAnnotation.user))\
                               .filter(TextAnnotation.document_id == document_id)
    
    # Applica filtro per categoria se specificato
    if selected_categories:
        annotations_query = annotations_query.join(Label).join(Category)\
                                           .filter(Category.name.in_(selected_categories))
    
    annotations = annotations_query.order_by(TextAnnotation.start_position).all()
    
    # Carica tutte le etichette per statistiche
    labels = Label.query.filter_by(is_active=True).order_by(Label.name).all()
    
    # Carica tutte le categorie per il filtro
    categories = Category.query.filter_by(is_active=True).order_by(Category.name).all()
    
    # Statistiche
    stats = {
        'total_annotations': len(annotations),
        'unique_labels': len(set(ann.label_id for ann in annotations)),
        'annotators': len(set(ann.user_id for ann in annotations)),
        'coverage_percentage': 0
    }
    
    # Calcola copertura
    if document.content and len(document.content) > 0:
        annotated_chars = sum(ann.end_position - ann.start_position for ann in annotations)
        stats['coverage_percentage'] = min(100, (annotated_chars / len(document.content)) * 100)
    
    return render_template('text_documents/document_annotations.html',
                         document=document,
                         annotations=annotations,
                         labels=labels,
                         categories=categories,
                         selected_categories=selected_categories,
                         stats=stats)

@text_documents_bp.route('/delete/<int:document_id>', methods=['POST'])
@login_required
def delete_document(document_id):
    """Elimina un documento e tutte le sue annotazioni"""
    document = TextDocument.query.get_or_404(document_id)
    
    # Verifica permessi
    if not current_user.is_admin and document.user_id != current_user.id:
        flash('Non hai i permessi per eliminare questo documento', 'error')
        return redirect(url_for('text_documents.list_documents'))
    
    try:
        # I file sono salvati nel filesystem ma non tracciamo il path nel DB
        # TODO: Implementare tracciamento path se necessario
        
        # Elimina annotazioni associate
        TextAnnotation.query.filter_by(document_id=document_id).delete()
        
        # Elimina documento
        db.session.delete(document)
        db.session.commit()
        
        flash('Documento eliminato con successo', 'success')
        
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Errore nell'eliminazione documento: {str(e)}")
        flash('Errore nell\'eliminazione del documento', 'error')
    
    return redirect(url_for('text_documents.list_documents'))

@text_documents_bp.route('/api/labels')
@login_required
def api_get_labels():
    """API per ottenere le etichette disponibili"""
    labels = Label.query.filter_by(is_active=True).order_by(Label.name).all()
    return jsonify({
        'success': True,
        'labels': [
            {
                'id': label.id,
                'name': label.name,
                'color': label.color,
                'description': label.description
            }
            for label in labels
        ]
    })

@text_documents_bp.route('/api/annotations/<int:annotation_id>', methods=['PUT'])
@login_required
def api_update_annotation(annotation_id):
    """API per modificare un'annotazione"""
    try:
        annotation = TextAnnotation.query.get_or_404(annotation_id)
        
        # Verifica permessi
        if not current_user.is_admin and annotation.user_id != current_user.id:
            return jsonify({'success': False, 'message': 'Non hai i permessi per modificare questa annotazione'}), 403
        
        data = request.get_json()
        
        # Aggiorna l'annotazione
        if 'label_id' in data:
            label = Label.query.get(data['label_id'])
            if not label:
                return jsonify({'success': False, 'message': 'Etichetta non valida'}), 400
            annotation.label_id = data['label_id']
        
        if 'note' in data:
            annotation.note = data['note'].strip() if data['note'] else None
        
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': 'Annotazione aggiornata con successo'
        })
        
    except Exception as e:
        current_app.logger.error(f"Errore aggiornamento annotazione: {str(e)}")
        db.session.rollback()
        return jsonify({'success': False, 'message': 'Errore interno del server'}), 500

@text_documents_bp.route('/api/annotations/<int:annotation_id>', methods=['DELETE'])
@login_required
def api_delete_annotation(annotation_id):
    """API per eliminare un'annotazione"""
    try:
        annotation = TextAnnotation.query.get_or_404(annotation_id)
        
        # Verifica permessi
        if not current_user.is_admin and annotation.user_id != current_user.id:
            return jsonify({'success': False, 'message': 'Non hai i permessi per eliminare questa annotazione'}), 403
        
        db.session.delete(annotation)
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': 'Annotazione eliminata con successo'
        })
        
    except Exception as e:
        current_app.logger.error(f"Errore eliminazione annotazione: {str(e)}")
        db.session.rollback()
        return jsonify({'success': False, 'message': 'Errore interno del server'}), 500

@text_documents_bp.route('/export/<int:document_id>/<format_type>')
@login_required
def export_annotations(document_id, format_type):
    """Esporta le annotazioni di un documento nei formati PDF, Word o LaTeX"""
    try:
        from services.annotation_export import AnnotationExporter
        
        document = TextDocument.query.get_or_404(document_id)
        
        # Verifica permessi
        if not current_user.is_admin and document.user_id != current_user.id:
            flash('Non hai i permessi per visualizzare questo documento', 'error')
            return redirect(url_for('text_documents.list_documents'))
        
        # Parametri di filtro
        selected_categories = request.args.getlist('categories')
        
        # Query per le annotazioni con filtro categoria
        annotations_query = db.session.query(TextAnnotation)\
                                   .options(db.joinedload(TextAnnotation.label).joinedload(Label.category_obj))\
                                   .options(db.joinedload(TextAnnotation.user))\
                                   .filter(TextAnnotation.document_id == document_id)
        
        # Applica filtro per categoria se specificato
        if selected_categories:
            annotations_query = annotations_query.join(Label).join(Category)\
                                               .filter(Category.name.in_(selected_categories))
        
        annotations = annotations_query.order_by(TextAnnotation.start_position).all()
        
        # Crea l'esportatore
        exporter = AnnotationExporter(document, annotations, selected_categories)
        
        # Esegui l'esportazione nel formato richiesto
        if format_type.lower() == 'pdf':
            return exporter.export_to_pdf()
        elif format_type.lower() == 'word':
            return exporter.export_to_word()
        elif format_type.lower() == 'latex':
            return exporter.export_to_latex()
        else:
            flash('Formato di esportazione non supportato', 'error')
            return redirect(url_for('text_documents.annotations', document_id=document_id))
            
    except ImportError:
        flash('Modulo di esportazione non disponibile. Installare le dipendenze necessarie.', 'error')
        return redirect(url_for('text_documents.annotations', document_id=document_id))
    except Exception as e:
        current_app.logger.error(f"Errore nell'esportazione: {str(e)}")
        flash('Errore durante l\'esportazione del documento', 'error')
        return redirect(url_for('text_documents.annotations', document_id=document_id))
