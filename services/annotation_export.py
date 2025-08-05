"""
Servizio per l'esportazione delle annotazioni in diversi formati
"""

import io
from datetime import datetime
from flask import make_response
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn

class AnnotationExporter:
    """Classe per l'esportazione delle annotazioni"""
    
    def __init__(self, document, annotations, filtered_categories=None):
        """
        Inizializza l'esportatore
        
        Args:
            document: Il documento di testo
            annotations: Lista delle annotazioni
            filtered_categories: Lista delle categorie da includere (None = tutte)
        """
        self.document = document
        self.annotations = annotations
        self.filtered_categories = filtered_categories or []
        
        # Filtra le annotazioni per categoria se specificato
        if self.filtered_categories:
            self.annotations = [
                ann for ann in self.annotations 
                if ann.label and ann.label.category_obj and ann.label.category_obj.name in self.filtered_categories
            ]
    
    def export_to_pdf(self):
        """
        Esporta le annotazioni in formato PDF
        
        Returns:
            Response: Response Flask con il PDF
        """
        buffer = io.BytesIO()
        
        # Crea il documento PDF
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18
        )
        
        # Stili
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=20,
            alignment=TA_CENTER
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=12,
            spaceAfter=10,
            textColor=colors.blue
        )
        
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=9,
            leading=11
        )
        
        # Contenuto del documento
        story = []
        
        # Titolo
        title = f"Annotazioni - {self.document.original_name}"
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 12))
        
        # Informazioni del documento
        doc_info = [
            ['Nome file:', self.document.original_name],
            ['Tipo documento:', self.document.document_type.title()],
            ['Data creazione:', self.document.created_at.strftime('%d/%m/%Y %H:%M')],
            ['Totale annotazioni:', str(len(self.annotations))],
            ['Data esportazione:', datetime.now().strftime('%d/%m/%Y %H:%M')]
        ]
        
        if self.filtered_categories:
            doc_info.append(['Categorie filtrate:', ', '.join(self.filtered_categories)])
        
        info_table = Table(doc_info, colWidths=[2*inch, 4*inch])
        info_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('ALIGN', (0, 0), (0, -1), 'RIGHT'),
            ('ALIGN', (1, 0), (1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('BOX', (0, 0), (-1, -1), 0.25, colors.black),
            ('GRID', (0, 0), (-1, -1), 0.25, colors.black),
        ]))
        
        story.append(info_table)
        story.append(Spacer(1, 20))
        
        # Statistiche per categoria
        if self.annotations:
            story.append(Paragraph("Statistiche per Categoria", heading_style))
            
            category_stats = {}
            for ann in self.annotations:
                if ann.label and ann.label.category_obj:
                    cat_name = ann.label.category_obj.name
                else:
                    cat_name = 'Senza categoria'
                
                if cat_name not in category_stats:
                    category_stats[cat_name] = 0
                category_stats[cat_name] += 1
            
            stats_data = [['Categoria', 'Numero Annotazioni']]
            for cat, count in sorted(category_stats.items()):
                stats_data.append([cat, str(count)])
            
            stats_table = Table(stats_data, colWidths=[3*inch, 1.5*inch])
            stats_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('ALIGN', (1, 0), (1, -1), 'CENTER'),
                ('BOX', (0, 0), (-1, -1), 0.25, colors.black),
                ('GRID', (0, 0), (-1, -1), 0.25, colors.black),
            ]))
            
            story.append(stats_table)
            story.append(Spacer(1, 20))
        
        # Lista delle annotazioni
        story.append(Paragraph("Elenco Annotazioni", heading_style))
        
        if self.annotations:
            # Raggruppa per categoria
            annotations_by_category = {}
            for ann in self.annotations:
                cat_name = ann.label.category_obj.name if ann.label.category_obj else 'Senza categoria'
                if cat_name not in annotations_by_category:
                    annotations_by_category[cat_name] = []
                annotations_by_category[cat_name].append(ann)
            
            for cat_name in sorted(annotations_by_category.keys()):
                # Intestazione categoria
                cat_style = ParagraphStyle(
                    'CategoryHeading',
                    parent=styles['Heading3'],
                    fontSize=12,
                    spaceAfter=8,
                    textColor=colors.darkblue
                )
                story.append(Paragraph(f"Categoria: {cat_name}", cat_style))
                
                # Annotazioni della categoria
                for i, ann in enumerate(annotations_by_category[cat_name], 1):
                    # Testo annotato COMPLETO - non tagliare mai
                    ann_text = ann.text_selection
                    
                    ann_data = [
                        [f"#{i}", ann.label.name, ann_text],
                        ['Posizione:', f"{ann.start_position}-{ann.end_position}", f"Utente: {ann.user.username}"],
                        ['Data:', ann.created_at.strftime('%d/%m/%Y %H:%M'), '']
                    ]
                    
                    ann_table = Table(ann_data, colWidths=[0.6*inch, 1.2*inch, 4.2*inch])
                    ann_table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (0, 0), colors.lightblue),
                        ('BACKGROUND', (1, 0), (1, 0), colors.lightgreen),
                        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
                        ('FONTSIZE', (0, 0), (-1, -1), 8),
                        ('ALIGN', (0, 0), (0, -1), 'CENTER'),
                        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                        ('BOX', (0, 0), (-1, -1), 0.25, colors.black),
                        ('GRID', (0, 0), (-1, -1), 0.25, colors.black),
                    ]))
                    
                    story.append(ann_table)
                    story.append(Spacer(1, 8))
                
                story.append(Spacer(1, 12))
        else:
            story.append(Paragraph("Nessuna annotazione trovata con i filtri specificati.", normal_style))
        
        # Genera il PDF
        doc.build(story)
        
        # Prepara la response
        buffer.seek(0)
        filename = f"annotazioni_{self.document.original_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        
        response = make_response(buffer.getvalue())
        response.headers['Content-Type'] = 'application/pdf'
        response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        return response
    
    def export_to_word(self):
        """
        Esporta le annotazioni in formato Word
        
        Returns:
            Response: Response Flask con il documento Word
        """
        doc = Document()
        
        # Titolo
        title = doc.add_heading(f'Annotazioni - {self.document.original_name}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Informazioni del documento
        doc.add_heading('Informazioni Documento', level=1)
        
        info_table = doc.add_table(rows=5 + (1 if self.filtered_categories else 0), cols=2)
        info_table.style = 'Table Grid'
        
        info_data = [
            ('Nome file:', self.document.original_name),
            ('Tipo documento:', self.document.document_type.title()),
            ('Data creazione:', self.document.created_at.strftime('%d/%m/%Y %H:%M')),
            ('Totale annotazioni:', str(len(self.annotations))),
            ('Data esportazione:', datetime.now().strftime('%d/%m/%Y %H:%M'))
        ]
        
        if self.filtered_categories:
            info_data.append(('Categorie filtrate:', ', '.join(self.filtered_categories)))
        
        for i, (label, value) in enumerate(info_data):
            row = info_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            # Formatta la prima colonna
            row.cells[0].paragraphs[0].runs[0].bold = True
        
        doc.add_page_break()
        
        # Statistiche per categoria
        if self.annotations:
            doc.add_heading('Statistiche per Categoria', level=1)
            
            category_stats = {}
            for ann in self.annotations:
                cat_name = ann.label.category_obj.name if ann.label.category_obj else 'Senza categoria'
                if cat_name not in category_stats:
                    category_stats[cat_name] = 0
                category_stats[cat_name] += 1
            
            stats_table = doc.add_table(rows=len(category_stats) + 1, cols=2)
            stats_table.style = 'Table Grid'
            
            # Header
            header_row = stats_table.rows[0]
            header_row.cells[0].text = 'Categoria'
            header_row.cells[1].text = 'Numero Annotazioni'
            for cell in header_row.cells:
                cell.paragraphs[0].runs[0].bold = True
            
            # Dati
            for i, (cat, count) in enumerate(sorted(category_stats.items()), 1):
                row = stats_table.rows[i]
                row.cells[0].text = cat
                row.cells[1].text = str(count)
        
        # Lista delle annotazioni
        doc.add_heading('Elenco Annotazioni', level=1)
        
        if self.annotations:
            # Raggruppa per categoria
            annotations_by_category = {}
            for ann in self.annotations:
                cat_name = ann.label.category_obj.name if ann.label.category_obj else 'Senza categoria'
                if cat_name not in annotations_by_category:
                    annotations_by_category[cat_name] = []
                annotations_by_category[cat_name].append(ann)
            
            for cat_name in sorted(annotations_by_category.keys()):
                # Intestazione categoria
                doc.add_heading(f'Categoria: {cat_name}', level=2)
                
                # Tabella per le annotazioni della categoria
                ann_table = doc.add_table(rows=len(annotations_by_category[cat_name]) + 1, cols=5)
                ann_table.style = 'Table Grid'
                
                # Header
                header_row = ann_table.rows[0]
                headers = ['#', 'Etichetta', 'Testo Annotato', 'Posizione', 'Utente/Data']
                for i, header in enumerate(headers):
                    header_row.cells[i].text = header
                    header_row.cells[i].paragraphs[0].runs[0].bold = True
                
                # Dati
                for i, ann in enumerate(annotations_by_category[cat_name], 1):
                    row = ann_table.rows[i]
                    row.cells[0].text = str(i)
                    row.cells[1].text = ann.label.name
                    
                    # Testo annotato COMPLETO - non tagliare mai
                    row.cells[2].text = ann.text_selection
                    
                    row.cells[3].text = f"{ann.start_position}-{ann.end_position}"
                    row.cells[4].text = f"{ann.user.username}\n{ann.created_at.strftime('%d/%m/%Y %H:%M')}"
        else:
            doc.add_paragraph("Nessuna annotazione trovata con i filtri specificati.")
        
        # Salva in buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Prepara la response
        filename = f"annotazioni_{self.document.original_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        response = make_response(buffer.getvalue())
        response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        return response
    
    def export_to_latex(self):
        """
        Esporta le annotazioni in formato LaTeX
        
        Returns:
            Response: Response Flask con il file LaTeX
        """
        latex_content = self._generate_latex_content()
        
        # Prepara la response
        filename = f"annotazioni_{self.document.original_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.tex"
        
        response = make_response(latex_content)
        response.headers['Content-Type'] = 'application/x-tex'
        response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        return response
    
    def _generate_latex_content(self):
        """
        Genera il contenuto LaTeX per le annotazioni
        
        Returns:
            str: Contenuto LaTeX
        """
        latex_lines = [
            "\\documentclass[a4paper,11pt]{article}",
            "\\usepackage[utf8]{inputenc}",
            "\\usepackage[T1]{fontenc}",
            "\\usepackage[italian]{babel}",
            "\\usepackage{geometry}",
            "\\usepackage{longtable}",
            "\\usepackage{booktabs}",
            "\\usepackage{xcolor}",
            "\\usepackage{fancyhdr}",
            "\\usepackage{hyperref}",
            "",
            "\\geometry{margin=2cm}",
            "\\pagestyle{fancy}",
            "\\fancyhf{}",
            "\\fancyhead[C]{Annotazioni - " + self._escape_latex(self.document.original_name) + "}",
            "\\fancyfoot[C]{\\thepage}",
            "",
            "\\title{Annotazioni - " + self._escape_latex(self.document.original_name) + "}",
            "\\author{Anatema - Sistema di Etichettatura}",
            f"\\date{{{datetime.now().strftime('%d/%m/%Y')}}}",
            "",
            "\\begin{document}",
            "",
            "\\maketitle",
            "",
            "\\section{Informazioni Documento}",
            "",
            "\\begin{tabular}{ll}",
            "\\textbf{Nome file:} & " + self._escape_latex(self.document.original_name) + " \\\\",
            "\\textbf{Tipo documento:} & " + self._escape_latex(self.document.document_type.title()) + " \\\\",
            "\\textbf{Data creazione:} & " + self.document.created_at.strftime('%d/%m/%Y %H:%M') + " \\\\",
            "\\textbf{Totale annotazioni:} & " + str(len(self.annotations)) + " \\\\",
            "\\textbf{Data esportazione:} & " + datetime.now().strftime('%d/%m/%Y %H:%M') + " \\\\",
        ]
        
        if self.filtered_categories:
            latex_lines.append("\\textbf{Categorie filtrate:} & " + self._escape_latex(', '.join(self.filtered_categories)) + " \\\\")
        
        latex_lines.extend([
            "\\end{tabular}",
            "",
            "\\section{Statistiche per Categoria}",
            ""
        ])
        
        if self.annotations:
            # Statistiche per categoria
            category_stats = {}
            for ann in self.annotations:
                cat_name = ann.label.category_obj.name if ann.label.category_obj else 'Senza categoria'
                if cat_name not in category_stats:
                    category_stats[cat_name] = 0
                category_stats[cat_name] += 1
            
            latex_lines.extend([
                "\\begin{tabular}{lr}",
                "\\toprule",
                "\\textbf{Categoria} & \\textbf{Numero Annotazioni} \\\\",
                "\\midrule"
            ])
            
            for cat, count in sorted(category_stats.items()):
                latex_lines.append(f"{self._escape_latex(cat)} & {count} \\\\")
            
            latex_lines.extend([
                "\\bottomrule",
                "\\end{tabular}",
                ""
            ])
            
            # Lista delle annotazioni
            latex_lines.extend([
                "\\section{Elenco Annotazioni}",
                ""
            ])
            
            # Raggruppa per categoria
            annotations_by_category = {}
            for ann in self.annotations:
                cat_name = ann.label.category_obj.name if ann.label.category_obj else 'Senza categoria'
                if cat_name not in annotations_by_category:
                    annotations_by_category[cat_name] = []
                annotations_by_category[cat_name].append(ann)
            
            for cat_name in sorted(annotations_by_category.keys()):
                latex_lines.extend([
                    f"\\subsection{{{self._escape_latex(cat_name)}}}",
                    "",
                    "\\begin{longtable}{p{0.8cm}p{2.5cm}p{8cm}p{1.5cm}p{2.2cm}}",
                    "\\toprule",
                    "\\textbf{\\#} & \\textbf{Etichetta} & \\textbf{Testo} & \\textbf{Posizione} & \\textbf{Utente/Data} \\\\",
                    "\\midrule",
                    "\\endfirsthead",
                    "",
                    "\\multicolumn{5}{c}%",
                    "{\\bfseries \\tablename\\ \\thetable{} -- continua dalla pagina precedente} \\\\",
                    "\\toprule",
                    "\\textbf{\\#} & \\textbf{Etichetta} & \\textbf{Testo} & \\textbf{Posizione} & \\textbf{Utente/Data} \\\\",
                    "\\midrule",
                    "\\endhead",
                    "",
                    "\\midrule \\multicolumn{5}{r}{{Continua nella pagina successiva}} \\\\ \\midrule",
                    "\\endfoot",
                    "",
                    "\\bottomrule",
                    "\\endlastfoot",
                    ""
                ])
                
                for i, ann in enumerate(annotations_by_category[cat_name], 1):
                    # Testo annotato COMPLETO - non tagliare mai
                    ann_text = self._escape_latex(ann.text_selection)
                    
                    latex_lines.append(
                        f"{i} & {self._escape_latex(ann.label.name)} & {ann_text} & "
                        f"{ann.start_position}-{ann.end_position} & "
                        f"{self._escape_latex(ann.user.username)} {ann.created_at.strftime('%d/%m/%Y')} \\\\"
                    )
                
                latex_lines.extend([
                    "\\end{longtable}",
                    ""
                ])
        else:
            latex_lines.append("Nessuna annotazione trovata con i filtri specificati.")
        
        latex_lines.extend([
            "",
            "\\end{document}"
        ])
        
        return "\n".join(latex_lines)
    
    def _escape_latex(self, text):
        """
        Escapa i caratteri speciali LaTeX
        
        Args:
            text (str): Testo da escape
            
        Returns:
            str: Testo con escape LaTeX
        """
        # Caratteri speciali LaTeX da escape
        latex_special_chars = {
            '&': '\\&',
            '%': '\\%',
            '$': '\\$',
            '#': '\\#',
            '^': '\\textasciicircum{}',
            '_': '\\_',
            '{': '\\{',
            '}': '\\}',
            '~': '\\textasciitilde{}',
            '\\': '\\textbackslash{}'
        }
        
        result = text
        for char, escape in latex_special_chars.items():
            result = result.replace(char, escape)
        
        return result
